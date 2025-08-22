# app.py
import os
import io
import json
import base64
from datetime import datetime
import psycopg2
from psycopg2.extras import DictCursor
from flask import (Flask, render_template, request, redirect, url_for,
                   flash, session, jsonify, Response, send_file, send_from_directory)
from flask_bcrypt import Bcrypt
from dotenv import load_dotenv
from functools import wraps
import magic
import uuid
import requests

# --- IMPORTAÇÃO PARA A API DO GOOGLE ---
from googleapiclient.discovery import build

import google.generativeai as genai

from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Inches

load_dotenv()

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
app.secret_key = os.environ.get('SECRET_KEY', 'uma-chave-secreta-forte-para-desenvolvimento')

UPLOAD_FOLDER = os.path.join(app.root_path, 'static', 'uploads')
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

bcrypt = Bcrypt(app)

# --- CONFIGURAÇÃO DAS APIs ---
genai.configure(api_key=os.environ.get("GEMINI_API_KEY"))
GOOGLE_API_KEY = os.environ.get("GOOGLE_SEARCH_API_KEY")
SEARCH_ENGINE_ID = os.environ.get("SEARCH_ENGINE_ID")

generation_config = {
    "temperature": 0.7,
    "top_p": 1,
    "top_k": 1,
    "max_output_tokens": 8192,
}
# --- CORREÇÃO DO NOME DO MODELO APLICADA AQUI ---
model = genai.GenerativeModel(model_name="gemini-2.5-flash", generation_config=generation_config)


def custom_search_images(query):
    """Busca por imagens usando a Custom Search JSON API."""
    if not GOOGLE_API_KEY or not SEARCH_ENGINE_ID:
        print("AVISO: Chave da API do Google ou ID do Motor de Busca não configurados.")
        return []
    try:
        service = build("customsearch", "v1", developerKey=GOOGLE_API_KEY)
        res = service.cse().list(
            q=query,
            cx=SEARCH_ENGINE_ID,
            searchType='image',
            num=5
        ).execute()
        return res.get('items', [])
    except Exception as e:
        print(f"Erro ao chamar a Custom Search API: {e}")
        return []


def get_db_connection():
    """Estabelece uma conexão com o banco de dados PostgreSQL."""
    try:
        conn_str = os.environ.get('DATABASE_URL')
        if conn_str:
            conn = psycopg2.connect(conn_str)
        else:
            conn = psycopg2.connect(
                host=os.environ.get('DB_HOST'),
                dbname=os.environ.get('DB_NAME'),
                user=os.environ.get('DB_USER'),
                password=os.environ.get('DB_PASSWORD'),
                port=os.environ.get('DB_PORT', 5432)
            )
        return conn
    except psycopg2.Error as e:
        print(f"Erro ao conectar ao banco de dados PostgreSQL: {e}")
        raise e


def clean_and_parse_json(response_text):
    """Limpa e tenta decodificar uma string JSON da resposta da IA."""
    if not response_text:
        return None
    clean_text = response_text.strip().replace('```json', '').replace('```', '').strip()
    try:
        return json.loads(clean_text)
    except json.JSONDecodeError as e:
        print(f"Erro ao decodificar JSON da IA: {e}")
        print(f"Texto recebido: {clean_text}")
        return None


@app.route('/uploads/<filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)


def search_questions_in_db(query_term):
    """Busca questões no banco de dados pelo termo fornecido."""
    conn = get_db_connection()
    cursor = conn.cursor(cursor_factory=DictCursor)
    like_term = f"%{query_term}%"
    cursor.execute("SELECT id, enunciado FROM questoes WHERE is_active = TRUE AND enunciado ILIKE %s LIMIT 10",
                   (like_term,))
    results = cursor.fetchall()
    cursor.close()
    conn.close()
    return results


def insert_question_in_db(question_data):
    """Insere uma nova questão e suas opções no banco de dados."""
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        dificuldade_map = {'FACIL': 'Fácil', 'MEDIO': 'Médio', 'DIFICIL': 'Difícil', 'MUITO_DIFICIL': 'Muito Difícil'}
        nivel_dificuldade_ia = question_data.get('nivel_dificuldade', 'MEDIO').upper()
        nivel_dificuldade_db = dificuldade_map.get(nivel_dificuldade_ia, 'Médio')

        tipo_questao_ia = question_data.get('tipo_questao', 'ESCOLHA_UNICA').upper()
        valid_types = ['ESCOLHA_UNICA', 'MULTIPLA_ESCOLHA', 'DISCURSIVA']
        tipo_questao_db = tipo_questao_ia if tipo_questao_ia in valid_types else 'ESCOLHA_UNICA'

        imagem_dados = None
        imagem_path = question_data.get('imagem_path')
        if imagem_path and os.path.exists(imagem_path):
            with open(imagem_path, 'rb') as f:
                imagem_dados = f.read()

        sql_questao = """
                      INSERT INTO questoes (enunciado, tipo_questao, autor_id, nivel_dificuldade, grau_ensino,
                                            area_conhecimento, imagem_url)
                      VALUES (%s, %s, %s, %s, %s, %s, %s) RETURNING id
                      """
        cursor.execute(sql_questao, (
            question_data.get('enunciado'),
            tipo_questao_db,
            session['user_id'],
            nivel_dificuldade_db,
            question_data.get('grau_ensino'),
            question_data.get('area_conhecimento'),
            imagem_dados
        ))
        questao_id = cursor.fetchone()[0]

        if 'opcoes' in question_data and questao_id:
            for opcao in question_data['opcoes']:
                sql_opcao = "INSERT INTO opcoes (questao_id, texto_opcao, is_correta) VALUES (%s, %s, %s)"
                cursor.execute(sql_opcao, (questao_id, opcao.get('texto_opcao'), bool(opcao.get('is_correta'))))

        conn.commit()
        return questao_id
    except psycopg2.Error as e:
        print(f"Erro ao inserir questão via IA: {e}")
        conn.rollback()
        return None
    finally:
        if cursor: cursor.close()
        if conn: conn.close()


def get_user_data():
    """Busca os dados do usuário, mas NUNCA guarda a foto na sessão."""
    if 'user_id' not in session:
        return None, None
    nome_completo = f"{session.get('user_nome', '')} {session.get('user_sobrenome', '')}".strip()
    foto_perfil_url = None
    conn = None
    try:
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=DictCursor)
        cursor.execute('SELECT foto_perfil FROM usuarios WHERE id = %s', (session['user_id'],))
        user = cursor.fetchone()
        if user and user.get('foto_perfil'):
            data = user['foto_perfil']
            foto_perfil_url = data.decode('utf-8') if isinstance(data, (bytes, bytearray)) else data
    except psycopg2.Error as e:
        print(f"Erro ao buscar dados do usuário: {e}")
    finally:
        if conn:
            cursor.close()
            conn.close()
    return nome_completo, foto_perfil_url


def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)

    return decorated_function


@app.route('/api/chat', methods=['POST'])
@login_required
def chat_ia():
    data = request.get_json()
    user_message = data.get('message')
    user_nome = session.get('user_nome', 'usuário')
    pending_action = "Sim" if 'pending_question' in session else "Não"

    intent_prompt = f"""
    Analise a mensagem do usuário chamado '{user_nome}' para determinar a intenção. As intenções possíveis são: SEARCH, CREATE, INSERT, CHAT.
    - Se o usuário quer procurar, pesquisar ou buscar, a intenção é SEARCH.
    - Se o usuário quer criar ou gerar uma questão, a intenção é CREATE.
    - Se uma questão foi recém-criada (pending_action='Sim') e a mensagem do usuário é afirmativa (sim, pode, cadastre, confirme), a intenção é INSERT.
    - Caso contrário, a intenção é CHAT.
    Extraia o tópico/termo de busca se a intenção for SEARCH ou CREATE.
    Responda APENAS com um JSON.
    Exemplo: {{"intent": "CREATE", "topic": "sistema solar"}}
    pending_action: {pending_action}
    Mensagem do usuário: "{user_message}"
    """
    try:
        response = model.generate_content(intent_prompt)
        intent_data = clean_and_parse_json(response.text)
        if not intent_data:
            raise ValueError("A IA não retornou um JSON de intenção válido.")
        intent = intent_data.get("intent")
        topic = intent_data.get("topic")

        if intent == "CREATE":
            create_prompt = f"""
            Crie uma questão de múltipla escolha sobre o tópico "{topic}".
            Formate a resposta como um JSON válido com as chaves: "enunciado", "tipo_questao", "nivel_dificuldade", "grau_ensino", "area_conhecimento", e "opcoes".
            A chave "tipo_questao" DEVE ter o valor "ESCOLHA_UNICA".
            A chave "opcoes" deve ser uma lista de 4 objetos, cada um com "texto_opcao" e "is_correta". Apenas uma opção deve ser correta.
            Responda APENAS com o JSON.
            """
            response = model.generate_content(create_prompt)
            question_json = clean_and_parse_json(response.text)
            if not question_json:
                raise ValueError("A IA não retornou um JSON de questão válido.")

            imagem_path_servidor = None
            search_results = custom_search_images(f"ilustração didática {topic}")
            if search_results:
                for item in search_results:
                    image_url = item.get('link')
                    if not image_url: continue
                    try:
                        headers = {
                            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3'}
                        img_response = requests.get(image_url, stream=True, timeout=10, headers=headers)
                        img_response.raise_for_status()
                        extensao = os.path.splitext(image_url)[1].split('?')[0] or '.jpg'
                        if extensao.lower() not in ['.jpg', '.jpeg', '.png', '.gif']: extensao = '.jpg'
                        nome_ficheiro = f"{uuid.uuid4()}{extensao}"
                        imagem_path_servidor = os.path.join(app.config['UPLOAD_FOLDER'], nome_ficheiro)
                        with open(imagem_path_servidor, 'wb') as f:
                            for chunk in img_response.iter_content(chunk_size=8192):
                                f.write(chunk)
                        question_json['imagem_path'] = imagem_path_servidor
                        break
                    except requests.exceptions.RequestException as e:
                        print(f"Falha ao descarregar {image_url}: {e}")
                        imagem_path_servidor = None
                        continue

            session['pending_question'] = question_json
            message = f"Criei a seguinte questão sobre '{topic}', {user_nome}:\n\n"
            message += f"**Enunciado:** {question_json.get('enunciado', 'N/A')}\n"
            for i, opt in enumerate(question_json.get('opcoes', [])):
                message += f"{i + 1}. {opt.get('texto_opcao', 'N/A')}\n"

            if imagem_path_servidor:
                try:
                    with open(imagem_path_servidor, 'rb') as f:
                        imagem_bytes = f.read()
                    mime_type = magic.from_buffer(imagem_bytes, mime=True)
                    base64_string = base64.b64encode(imagem_bytes).decode('utf-8')
                    data_url = f"data:{mime_type};base64,{base64_string}"
                    message += f"\n<div class='image-preview-ia'><img src='{data_url}' alt='Imagem sugerida'></div>\n"
                except Exception as e:
                    print(f"Erro ao embutir a imagem: {e}")

            message += "\nVocê gostaria de cadastrá-la no banco de dados?"
            return jsonify({'type': 'chat', 'message': message})

        elif intent == "SEARCH":
            results = search_questions_in_db(topic)
            message = (f"Encontrei {len(results)} questões sobre '{topic}', {user_nome}:\n" +
                       "".join([f"- #{res['id']}: {res['enunciado'][:80]}...\n" for res in results])
                       if results else f"Não encontrei nenhuma questão sobre '{topic}'. Gostaria que eu criasse uma para você?")
            return jsonify({'type': 'chat', 'message': message})

        elif intent == "INSERT":
            pending_question = session.get('pending_question')
            if pending_question:
                new_question_id = insert_question_in_db(pending_question)
                message = (f"Perfeito, {user_nome}! A questão #{new_question_id} foi cadastrada com sucesso. ✅"
                           if new_question_id else "Ocorreu um erro ao tentar cadastrar a questão.")
                session.pop('pending_question', None)
            else:
                message = "Não encontrei nenhuma questão pendente para cadastrar."
            return jsonify({'type': 'chat', 'message': message})

        else:  # CHAT
            chat_prompt = f"Você é um assistente de IA amigável. O nome do usuário é {user_nome}. Responda à seguinte mensagem: \"{user_message}\""
            response = model.generate_content(chat_prompt)
            return jsonify({'type': 'chat', 'message': response.text})

    except Exception as e:
        print(f"Erro na API do Gemini ou no processamento do chat: {e}")
        session.pop('pending_question', None)
        return jsonify({'type': 'chat', 'message': 'Desculpe, ocorreu um erro. Poderia reformular seu pedido?'}), 500


@app.route('/')
def index():
    if 'user_id' in session:
        return redirect(url_for('painel'))
    return redirect(url_for('login'))


@app.route('/login', methods=['GET', 'POST'])
def login():
    if 'user_id' in session:
        return redirect(url_for('painel'))
    if request.method == 'POST':
        email = request.form.get('email')
        senha = request.form.get('senha')
        conn = None
        try:
            conn = get_db_connection()
            cursor = conn.cursor(cursor_factory=DictCursor)
            cursor.execute('SELECT id, nome, sobrenome, senha_hash, foto_perfil FROM usuarios WHERE email = %s',
                           (email,))
            user = cursor.fetchone()
            if user and bcrypt.check_password_hash(user['senha_hash'], senha):
                session.clear()
                session['user_id'] = user['id']
                session['user_nome'] = user['nome']
                session['user_sobrenome'] = user['sobrenome']

                foto_perfil_data = user.get('foto_perfil')
                foto_perfil_url = None
                if foto_perfil_data:
                    foto_perfil_url = foto_perfil_data.decode('utf-8') if isinstance(foto_perfil_data, (bytes,
                                                                                                        bytearray)) else foto_perfil_data

                return jsonify({
                    'success': True,
                    'user': {
                        'nome_completo': f"{user['nome']} {user['sobrenome']}".strip(),
                        'foto_perfil_url': foto_perfil_url
                    },
                    'redirect_url': url_for('painel')
                })
            else:
                return jsonify({'success': False, 'message': 'Email ou senha inválidos.'}), 401
        except psycopg2.Error as e:
            print(f"Erro no login: {e}")
            return jsonify({'success': False, 'message': 'Erro ao conectar com o banco de dados.'}), 500
        finally:
            if conn:
                cursor.close()
                conn.close()
    return render_template('login.html')


@app.route('/painel')
@login_required
def painel():
    nome_completo, foto_perfil_url = get_user_data()
    return render_template('painel.html', nome_completo=nome_completo, foto_perfil_url=foto_perfil_url, view='home')


@app.route('/search_questoes')
@login_required
def search_questoes():
    query = request.args.get('q', '')
    if not query or len(query) < 2:
        return jsonify([])
    conn = None
    try:
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=DictCursor)
        sql_search = """
                     SELECT id, enunciado, tipo_questao, nivel_dificuldade, grau_ensino, area_conhecimento
                     FROM questoes
                     WHERE is_active = TRUE
                       AND (enunciado ILIKE %s OR nivel_dificuldade::text ILIKE %s OR grau_ensino ILIKE %s OR area_conhecimento ILIKE %s)
                     ORDER BY id DESC LIMIT 10
                     """
        like_term = f"%{query}%"
        cursor.execute(sql_search, (like_term, like_term, like_term, like_term))
        results = [dict(row) for row in cursor.fetchall()]
        return jsonify(results)
    except psycopg2.Error as e:
        print(f"Erro na busca de questões: {e}")
        return jsonify({'error': 'Erro no servidor'}), 500
    finally:
        if conn:
            cursor.close()
            conn.close()


@app.route('/banco_questoes')
@login_required
def banco_questoes():
    nome_completo, foto_perfil_url = get_user_data()
    search_query = request.args.get('q', '')
    nivel_dificuldade = request.args.get('nivel', '')
    grau_ensino = request.args.get('grau', '')
    area_conhecimento = request.args.get('area', '')
    lista_questoes = []
    conn = None
    try:
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=DictCursor)
        sql_query = "SELECT id, enunciado, tipo_questao, nivel_dificuldade, grau_ensino, area_conhecimento FROM questoes WHERE is_active = TRUE"
        params = []
        if search_query:
            sql_query += " AND enunciado ILIKE %s"
            params.append(f"%{search_query}%")
        if nivel_dificuldade:
            sql_query += " AND nivel_dificuldade = %s"
            params.append(nivel_dificuldade)
        if grau_ensino:
            sql_query += " AND grau_ensino = %s"
            params.append(grau_ensino)
        if area_conhecimento:
            sql_query += " AND area_conhecimento ILIKE %s"
            params.append(f"%{area_conhecimento}%")
        sql_query += " ORDER BY id DESC"
        cursor.execute(sql_query, tuple(params))
        lista_questoes = cursor.fetchall()
    except psycopg2.Error as e:
        flash("Erro ao carregar as questões.", "error")
        print(f"Erro em /banco_questoes: {e}")
    finally:
        if conn:
            cursor.close()
            conn.close()
    return render_template('painel.html',
                           nome_completo=nome_completo,
                           foto_perfil_url=foto_perfil_url,
                           view='banco_questoes',
                           questoes=lista_questoes,
                           search_query=search_query,
                           nivel_dificuldade=nivel_dificuldade,
                           grau_ensino=grau_ensino,
                           area_conhecimento=area_conhecimento)


@app.route('/cadastrar_questoes')
@login_required
def cadastrar_questoes():
    nome_completo, foto_perfil_url = get_user_data()
    return render_template('painel.html', nome_completo=nome_completo, foto_perfil_url=foto_perfil_url,
                           view='cadastrar_questoes')


@app.route('/chat_ia')
@login_required
def chat_page():
    nome_completo, foto_perfil_url = get_user_data()
    user_nome = session.get('user_nome', '')
    return render_template('painel.html',
                           nome_completo=nome_completo,
                           foto_perfil_url=foto_perfil_url,
                           view='chat_ia',
                           user_nome=user_nome)


@app.route('/generate_questao', methods=['POST'])
@login_required
def generate_questao():
    data = request.get_json()
    tipo = data.get('tipo', 'ESCOLHA_UNICA')
    nivel = data.get('nivel', 'Fácil')
    grau = data.get('grau', 'Ensino Fundamental')
    area = data.get('area', 'Conhecimentos Gerais')
    prompt = (f"Gere uma questão de {tipo.replace('_', ' ').lower()} "
              f"com dificuldade {nivel.lower()} para o {grau.lower()} sobre {area}. "
              "O enunciado deve ser claro. Para questões de escolha única ou múltipla, "
              "gere 4 opções, e uma ou mais delas deve ser a correta. "
              "Formate a resposta como um JSON válido com as seguintes chaves: "
              "'enunciado', 'tipo_questao', 'nivel_dificuldade', 'grau_ensino', 'area_conhecimento', 'opcoes'. "
              "A chave 'opcoes' deve ser uma lista de objetos, cada um com as chaves 'texto_opcao' e 'is_correta' (booleano). "
              "Responda APENAS com o JSON.")
    try:
        response = model.generate_content(prompt)
        questao_gerada = clean_and_parse_json(response.text)
        if not questao_gerada:
            raise ValueError("A IA não retornou um JSON de questão válido.")
        return jsonify(questao_gerada)
    except Exception as e:
        print(f"Erro ao gerar questão com Gemini: {e}")
        return jsonify({'error': 'Falha ao gerar questão com a IA.'}), 500


@app.route('/lixeira')
@login_required
def lixeira():
    nome_completo, foto_perfil_url = get_user_data()
    lista_questoes_excluidas = []
    conn = None
    try:
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=DictCursor)
        cursor.execute("SELECT id, enunciado, tipo_questao FROM questoes WHERE is_active = FALSE ORDER BY id DESC")
        lista_questoes_excluidas = cursor.fetchall()
    except psycopg2.Error as e:
        flash("Erro ao carregar a lixeira.", "error")
        print(f"Erro em /lixeira: {e}")
    finally:
        if conn:
            cursor.close()
            conn.close()
    return render_template('painel.html', nome_completo=nome_completo, foto_perfil_url=foto_perfil_url, view='lixeira',
                           questoes=lista_questoes_excluidas)


@app.route('/configuracoes')
@login_required
def configuracoes():
    nome_completo, foto_perfil_url = get_user_data()
    user_data = {}
    conn = None
    try:
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=DictCursor)
        cursor.execute('SELECT nome, sobrenome, email FROM usuarios WHERE id = %s', (session['user_id'],))
        user_data = cursor.fetchone()
    except psycopg2.Error as e:
        flash("Não foi possível carregar os seus dados.", "error")
        print(f"Erro em /configuracoes: {e}")
    finally:
        if conn:
            cursor.close()
            conn.close()
    return render_template('painel.html', nome_completo=nome_completo, foto_perfil_url=foto_perfil_url,
                           view='configuracoes', user=user_data)


@app.route('/update_profile', methods=['POST'])
@login_required
def update_profile():
    nome = request.form.get('nome')
    sobrenome = request.form.get('sobrenome')
    if not nome or not sobrenome:
        flash("Nome e sobrenome são obrigatórios.", "error")
        return redirect(url_for('configuracoes'))
    conn = None
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("UPDATE usuarios SET nome = %s, sobrenome = %s WHERE id = %s",
                       (nome, sobrenome, session['user_id']))
        conn.commit()
        session['user_nome'] = nome
        session['user_sobrenome'] = sobrenome
        flash("Perfil atualizado com sucesso!", "success")
    except psycopg2.Error as e:
        if conn: conn.rollback()
        flash("Ocorreu um erro ao atualizar o seu perfil.", "error")
        print(f"Erro em /update_profile: {e}")
    finally:
        if conn:
            cursor.close()
            conn.close()
    return redirect(url_for('configuracoes'))


@app.route('/change_password', methods=['POST'])
@login_required
def change_password():
    senha_atual = request.form.get('senha_atual')
    nova_senha = request.form.get('nova_senha')
    confirmar_senha = request.form.get('confirmar_senha')
    if not all([senha_atual, nova_senha, confirmar_senha]):
        flash("Todos os campos de senha são obrigatórios.", "error")
        return redirect(url_for('configuracoes'))
    if nova_senha != confirmar_senha:
        flash("As novas senhas não coincidem.", "error")
        return redirect(url_for('configuracoes'))
    conn = None
    try:
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=DictCursor)
        cursor.execute("SELECT senha_hash FROM usuarios WHERE id = %s", (session['user_id'],))
        user = cursor.fetchone()
        if not user or not bcrypt.check_password_hash(user['senha_hash'], senha_atual):
            flash("A sua senha atual está incorreta.", "error")
            return redirect(url_for('configuracoes'))
        nova_senha_hash = bcrypt.generate_password_hash(nova_senha).decode('utf-8')
        cursor.execute("UPDATE usuarios SET senha_hash = %s WHERE id = %s", (nova_senha_hash, session['user_id']))
        conn.commit()
        flash("Senha alterada com sucesso!", "success")
    except psycopg2.Error as e:
        if conn: conn.rollback()
        flash("Ocorreu um erro ao alterar a sua senha.", "error")
        print(f"Erro em /change_password: {e}")
    finally:
        if conn:
            cursor.close()
            conn.close()
    return redirect(url_for('configuracoes'))


@app.route('/get_questao/<int:questao_id>')
@login_required
def get_questao(questao_id):
    conn = None
    try:
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=DictCursor)
        cursor.execute(
            "SELECT id, enunciado, tipo_questao, autor_id, nivel_dificuldade, grau_ensino, area_conhecimento, imagem_url FROM questoes WHERE id = %s",
            (questao_id,))
        questao = cursor.fetchone()
        if not questao:
            return jsonify({'error': 'Questão não encontrada'}), 404

        questao_dict = dict(questao)
        imagem_bytes = questao.get('imagem_url')

        if imagem_bytes:
            if isinstance(imagem_bytes, memoryview):
                imagem_bytes = bytes(imagem_bytes)
            try:
                mime_type = magic.from_buffer(imagem_bytes, mime=True)
                questao_dict['imagem_url'] = f"data:{mime_type};base64,{base64.b64encode(imagem_bytes).decode('utf-8')}"
            except Exception as e:
                print(f"Não foi possível processar a imagem da questão #{questao_id}: {e}")
                questao_dict['imagem_url'] = None
        else:
            questao_dict['imagem_url'] = None

        if questao['tipo_questao'] != 'DISCURSIVA':
            cursor.execute("SELECT texto_opcao, is_correta, imagem_url FROM opcoes WHERE questao_id = %s ORDER BY id",
                           (questao_id,))
            opcoes_raw = cursor.fetchall()
            opcoes_list = []
            for op in opcoes_raw:
                op_dict = dict(op)
                op_imagem_bytes = op.get('imagem_url')
                if op_imagem_bytes:
                    if isinstance(op_imagem_bytes, memoryview):
                        op_imagem_bytes = bytes(op_imagem_bytes)
                    try:
                        mime_type = magic.from_buffer(op_imagem_bytes, mime=True)
                        op_dict[
                            'imagem_url'] = f"data:{mime_type};base64,{base64.b64encode(op_imagem_bytes).decode('utf-8')}"
                    except Exception:
                        op_dict['imagem_url'] = None
                else:
                    op_dict['imagem_url'] = None
                opcoes_list.append(op_dict)
            questao_dict['opcoes'] = opcoes_list

        return jsonify(questao_dict)
    except psycopg2.Error as e:
        print(f"Erro em /get_questao: {e}")
        return jsonify({'error': 'Erro no servidor'}), 500
    finally:
        if conn:
            cursor.close()
            conn.close()


@app.route('/delete_questao/<int:questao_id>', methods=['POST'])
@login_required
def delete_questao(questao_id):
    conn = None
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("UPDATE questoes SET is_active = FALSE WHERE id = %s AND autor_id = %s",
                       (questao_id, session['user_id']))
        if cursor.rowcount == 0:
            conn.rollback()
            return jsonify({'success': False, 'error': 'Questão não encontrada ou sem permissão.'}), 404
        conn.commit()
        return jsonify({'success': True, 'message': 'Questão movida para a lixeira!'})
    except psycopg2.Error as e:
        if conn: conn.rollback()
        print(f"Erro em /delete_questao: {e}")
        return jsonify({'success': False, 'error': 'Erro no servidor.'}), 500
    finally:
        if conn:
            cursor.close()
            conn.close()


@app.route('/restore_questao/<int:questao_id>', methods=['POST'])
@login_required
def restore_questao(questao_id):
    conn = None
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("UPDATE questoes SET is_active = TRUE WHERE id = %s AND autor_id = %s",
                       (questao_id, session['user_id']))
        if cursor.rowcount == 0:
            conn.rollback()
            return jsonify({'success': False, 'error': 'Questão não encontrada ou sem permissão.'}), 404
        conn.commit()
        return jsonify({'success': True, 'message': 'Questão restaurada!'})
    except psycopg2.Error as e:
        if conn: conn.rollback()
        print(f"Erro em /restore_questao: {e}")
        return jsonify({'success': False, 'error': 'Erro no servidor.'}), 500
    finally:
        if conn:
            cursor.close()
            conn.close()


@app.route('/delete_permanently/<int:questao_id>', methods=['POST'])
@login_required
def delete_permanently(questao_id):
    conn = None
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("DELETE FROM questoes WHERE id = %s AND autor_id = %s", (questao_id, session['user_id']))
        if cursor.rowcount == 0:
            conn.rollback()
            return jsonify({'success': False, 'error': 'Questão não encontrada ou sem permissão.'}), 404
        conn.commit()
        return jsonify({'success': True, 'message': 'Questão excluída permanentemente!'})
    except psycopg2.Error as e:
        if conn: conn.rollback()
        print(f"Erro em /delete_permanently: {e}")
        return jsonify({'success': False, 'error': 'Erro no servidor.'}), 500
    finally:
        if conn:
            cursor.close()
            conn.close()


@app.route('/edit_questao/<int:questao_id>', methods=['POST'])
@login_required
def edit_questao(questao_id):
    enunciado = request.form.get('enunciado')
    nivel_dificuldade_form = request.form.get('nivel_dificuldade')
    grau_ensino = request.form.get('grau_ensino')
    area_conhecimento = request.form.get('area_conhecimento')
    imagem_questao = request.files.get('imagem')
    imagem_questao_dados = imagem_questao.read() if imagem_questao and imagem_questao.filename else None
    if not all([enunciado, nivel_dificuldade_form]):
        flash("Enunciado e Nível de Dificuldade são obrigatórios.", "error")
        return redirect(url_for('banco_questoes'))
    dificuldade_map = {'FACIL': 'Fácil', 'MEDIO': 'Médio', 'DIFICIL': 'Difícil', 'MUITO_DIFICIL': 'Muito Difícil'}
    nivel_dificuldade_db = dificuldade_map.get(nivel_dificuldade_form.upper().replace("_", " "), nivel_dificuldade_form)
    conn = None
    try:
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=DictCursor)
        cursor.execute("SELECT autor_id, tipo_questao FROM questoes WHERE id = %s", (questao_id,))
        result = cursor.fetchone()
        if not result or result['autor_id'] != session['user_id']:
            flash("Você não tem permissão para editar esta questão.", "error")
            return redirect(url_for('banco_questoes'))
        tipo_questao = result['tipo_questao']
        sql_update = """
                     UPDATE questoes \
                     SET enunciado         = %s, \
                         nivel_dificuldade = %s, \
                         grau_ensino       = %s, \
                         area_conhecimento = %s, \
                         imagem_url        = COALESCE(%s, imagem_url)
                     WHERE id = %s
                     """
        cursor.execute(sql_update,
                       (enunciado, nivel_dificuldade_db, grau_ensino, area_conhecimento, imagem_questao_dados,
                        questao_id))
        cursor.execute("DELETE FROM opcoes WHERE questao_id = %s", (questao_id,))
        if tipo_questao in ['ESCOLHA_UNICA', 'MULTIPLA_ESCOLHA']:
            opcoes_texto = request.form.getlist('opcoes_texto[]')
            opcoes_imagens = request.files.getlist('opcoes_imagem[]')
            respostas_corretas_indices = request.form.getlist('respostas_corretas[]')
            for i, texto_opcao in enumerate(opcoes_texto):
                imagem_opcao_dados = opcoes_imagens[i].read() if i < len(opcoes_imagens) and opcoes_imagens[
                    i].filename else None
                if not texto_opcao and not imagem_opcao_dados: continue
                is_correta = str(i) in respostas_corretas_indices
                sql_opcao = "INSERT INTO opcoes (questao_id, texto_opcao, is_correta, imagem_url) VALUES (%s, %s, %s, %s)"
                cursor.execute(sql_opcao, (questao_id, texto_opcao, is_correta, imagem_opcao_dados))
        conn.commit()
        flash("Questão atualizada com sucesso!", "success")
    except (psycopg2.Error, ValueError) as e:
        if conn: conn.rollback()
        flash(f"Erro ao atualizar a questão: {e}", "error")
        print(f"Erro em /edit_questao: {e}")
    finally:
        if conn:
            cursor.close()
            conn.close()
    return redirect(url_for('banco_questoes'))


@app.route('/add_questao', methods=['POST'])
@login_required
def add_questao():
    tipo_questao = request.form.get('tipo_questao')
    enunciado = request.form.get('enunciado')
    nivel_dificuldade_form = request.form.get('nivel_dificuldade')
    grau_ensino = request.form.get('grau_ensino')
    area_conhecimento = request.form.get('area_conhecimento')
    imagem_questao = request.files.get('imagem')
    imagem_questao_dados = imagem_questao.read() if imagem_questao and imagem_questao.filename else None
    if not all([tipo_questao, enunciado, nivel_dificuldade_form]):
        flash("Todos os campos principais são obrigatórios.", "error")
        return redirect(url_for('cadastrar_questoes'))
    dificuldade_map = {'FACIL': 'Fácil', 'MEDIO': 'Médio', 'DIFICIL': 'Difícil', 'MUITO_DIFICIL': 'Muito Difícil'}
    nivel_dificuldade_db = dificuldade_map.get(nivel_dificuldade_form.upper().replace("_", " "), nivel_dificuldade_form)
    conn = None
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        sql_questao = """
                      INSERT INTO questoes (enunciado, tipo_questao, autor_id, nivel_dificuldade, grau_ensino,
                                            area_conhecimento, imagem_url)
                      VALUES (%s, %s, %s, %s, %s, %s, %s) RETURNING id
                      """
        cursor.execute(sql_questao, (enunciado, tipo_questao, session['user_id'], nivel_dificuldade_db, grau_ensino,
                                     area_conhecimento, imagem_questao_dados))
        questao_id = cursor.fetchone()[0]
        if tipo_questao in ['ESCOLHA_UNICA', 'MULTIPLA_ESCOLHA']:
            opcoes_texto = request.form.getlist('opcoes_texto[]')
            opcoes_imagens = request.files.getlist('opcoes_imagem[]')
            respostas_corretas_indices = request.form.getlist('respostas_corretas[]')
            if not opcoes_texto:
                raise ValueError("Questões de múltipla escolha precisam de opções.")
            for i, texto_opcao in enumerate(opcoes_texto):
                imagem_opcao_dados = opcoes_imagens[i].read() if i < len(opcoes_imagens) and opcoes_imagens[
                    i].filename else None
                if not texto_opcao and not imagem_opcao_dados: continue
                is_correta = str(i) in respostas_corretas_indices
                sql_opcao = "INSERT INTO opcoes (questao_id, texto_opcao, is_correta, imagem_url) VALUES (%s, %s, %s, %s)"
                cursor.execute(sql_opcao, (questao_id, texto_opcao, is_correta, imagem_opcao_dados))
        conn.commit()
        flash("Questão cadastrada com sucesso!", "success")
    except (psycopg2.Error, ValueError) as e:
        if conn: conn.rollback()
        flash(f"Erro ao cadastrar a questão: {e}", "error")
        print(f"Erro em /add_questao: {e}")
        return redirect(url_for('cadastrar_questoes'))
    finally:
        if conn:
            cursor.close()
            conn.close()
    return redirect(url_for('banco_questoes'))


@app.route('/upload_foto', methods=['POST'])
@login_required
def upload_foto():
    image_data = request.get_json().get('image')
    if not image_data:
        return jsonify({'success': False, 'error': 'Dados da imagem em falta'}), 400
    conn = None
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("UPDATE usuarios SET foto_perfil = %s WHERE id = %s", (image_data, session['user_id']))
        conn.commit()
        return jsonify({'success': True})
    except psycopg2.Error as e:
        if conn: conn.rollback()
        print(f"Erro em /upload_foto: {e}")
        return jsonify({'success': False, 'error': 'Erro no servidor.'}), 500
    finally:
        if conn:
            cursor.close()
            conn.close()


@app.route('/logout')
def logout():
    session.clear()
    flash("Você saiu da sua conta.")
    return redirect(url_for('login'))


@app.route('/export_questoes', methods=['POST'])
@login_required
def export_questoes():
    """Exporta questões selecionadas como um arquivo .docx e o retorna como download.
    Espera JSON { "ids": [1,2,3] } ou um form com 'ids' como CSV ou 'ids[]'.
    """
    # Obter lista de ids do pedido (suporta JSON e form data)
    ids = []
    try:
        if request.is_json:
            payload = request.get_json() or {}
            ids = payload.get('ids', [])
        else:
            ids = request.form.getlist('ids[]') or request.form.get('ids')
            if isinstance(ids, str):
                ids = [i.strip() for i in ids.split(',') if i.strip()]
    except Exception:
        ids = []

    # Normalizar para inteiros
    try:
        ids = [int(i) for i in ids]
    except Exception:
        return jsonify({'error': 'IDs inválidos'}), 400

    if not ids:
        return jsonify({'error': 'Nenhum ID fornecido para exportação.'}), 400

    conn = None
    try:
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=DictCursor)
        cursor.execute(
            """SELECT id, enunciado, tipo_questao, nivel_dificuldade, grau_ensino, area_conhecimento, imagem_url
               FROM questoes WHERE id = ANY(%s) ORDER BY id""",
            (ids,)
        )
        questoes = cursor.fetchall()

        if not questoes:
            return jsonify({'error': 'Nenhuma questão encontrada para os IDs fornecidos.'}), 404

        doc = Document()
        doc.styles['Normal'].font.name = 'Arial'
        doc.styles['Normal'].font.size = Pt(11)

        for idx, q in enumerate(questoes, start=1):
            # Apenas adicionar o enunciado (sem número da questão) — não adicionar meta/linha de tipo, nível, grau ou área
            enunciado = q.get('enunciado') or ''
            doc.add_paragraph(enunciado)

            # Inserir imagem da questão, se houver
            imagem_bytes = q.get('imagem_url')
            if imagem_bytes:
                try:
                    if isinstance(imagem_bytes, memoryview):
                        imagem_bytes = bytes(imagem_bytes)
                    img_io = io.BytesIO(imagem_bytes)
                    # Adicionar imagem com largura máxima de 4 polegadas
                    doc.add_picture(img_io, width=Inches(4))
                except Exception as e:
                    print(f"Erro ao inserir imagem da questão {q['id']}: {e}")

            # Opções (sem marcações extras além das letras e texto/imagem)
            if q.get('tipo_questao') != 'DISCURSIVA':
                cursor.execute("SELECT texto_opcao, is_correta, imagem_url FROM opcoes WHERE questao_id = %s ORDER BY id",
                               (q['id'],))
                opcoes = cursor.fetchall()
                if opcoes:
                    p = doc.add_paragraph()
                    for i, op in enumerate(opcoes):
                        letra = chr(ord('A') + i)
                        texto = op.get('texto_opcao') or ''
                        p.add_run(f"{letra}. {texto}")
                        p.add_run('\n')
                        op_img = op.get('imagem_url')
                        if op_img:
                            try:
                                if isinstance(op_img, memoryview):
                                    op_img = bytes(op_img)
                                img_io = io.BytesIO(op_img)
                                doc.add_picture(img_io, width=Inches(3))
                            except Exception as e:
                                print(f"Erro ao inserir imagem da opção da questão {q['id']}: {e}")

            doc.add_page_break()

        # Preparar arquivo em memória
        out_io = io.BytesIO()
        doc.save(out_io)
        out_io.seek(0)

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"questoes_export_{timestamp}.docx"

        return send_file(out_io,
                         as_attachment=True,
                         download_name=filename,
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    except psycopg2.Error as e:
        print(f"Erro ao exportar questões: {e}")
        return jsonify({'error': 'Erro no servidor ao exportar.'}), 500
    finally:
        if conn:
            cursor.close()
            conn.close()


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 4000))
    app.run(host='0.0.0.0', port=port, debug=True)